# Michael Navarro - Program #13

#--------------- Libraries ------------------#
import PyPDF2
import docx
import csv
import datetime
import json

#--------------- Helper Functions ------------------#

#--------------- Requirement #1 ------------------#
'''Output a header in the console: “This is Program13 - <yournamehere>”'''
print('\nRequirement #1\n')
print('This is Program13 - Michael Navarro')

#--------------- Requirement #2 ------------------#
'''Print “This program demonstrates using Python to work with PDF 
and MS Word documents,and CSV and JSON data.”'''
print('\nRequirement #2\n')
print('This program demonstrates using Python to work with PDF and MS Word documents, and CSV and JSON data.')

#--------------- Requirement #3 ------------------#
'''Use classes and functions to organize the functionality of this program. 
You should have the following classes: PDFProcessing, WordProcessing, CSVProcessing, and JSONProcessing. 
Include the appropriate data and functions in each class to perform the requirements below.'''
print('\nRequirement #3\n')

class PDFProcessing:
    def __init__(self, PDF_filename):
        self.PDF_filename = PDF_filename
        self.PDF_Obj = open(PDF_filename, 'rb')
        self.PDF_data = PyPDF2.PdfFileReader(self.PDF_Obj)
    # End init

    def display_num_pages(self):
        print(f'The file "{self.PDF_filename}" has {self.PDF_data.getNumPages()} pages.')
    # End display_num_pages

    def display_page_data(self):
        page_num = int(input('Please enter the page number you would like to look at: ')) - 1

        PDF_page = self.PDF_data.getPage(page_num)
        print(PDF_page.extractText())
    # End display_page_data

    def close_pdf(self):
        self.PDF_Obj.close()
    # End close_pdf
# End PDFProcessing

class WordProcessing:
    def __init__(self, Word_filename):
        self.Word_filename = Word_filename
        self.Word_data = docx.Document(Word_filename)
    # End init

    def display_num_paragraphs(self):
        print(f'The file "{self.Word_filename}" has {len(self.Word_data.paragraphs)} paragraphs.')
    # End display_num_paragraphs

    def display_paragraph_text(self):
        paragraph_num = int(input('Please enter the paragraph number you would like to view: ')) - 1
        print(f'\n{self.Word_data.paragraphs[paragraph_num].text}')
    # End display_paragraph_text
# End WordProcessing

class CSVProcessing:
    def __init__(self, CSV_filename):
        self.CSV_filename = CSV_filename
        self.CSV_Obj = open(CSV_filename)
        self.CSV_data = csv.reader(self.CSV_Obj)
    # End init

    def display_contents(self):
        content_dict = csv.DictReader(self.CSV_Obj, ['time', 'item', 'amount'])
        for row in content_dict:
            print(row['time'], row['item'], row['amount'])
    # End display_contents

    def write_data(self):
        current_time = datetime.datetime.now()
        output_file = open(self.CSV_filename, 'a', newline='')
        output_handler = csv.writer(output_file)

        item = input('Please enter the item you would like to record: ').title()
        amount = input('Please enter the amount: ')

        output_handler.writerow([current_time.strftime("%m-%d-%y %H:%M"), item, amount])
        output_file.close()
    # End write_data

    def close_csv(self):
        self.CSV_Obj.close()
    # End close_csv
# End CSVProcessing

class JSONProcessing:
    def __init__(self):
        self.city_dict = {}
    # End init

    def update_dict(self):
        num_cities = int(input('Please enter the number of cities you wish to record: '))

        for num in range(num_cities):
            current_city_name = input('\nPlease enter the name of a city: ').title()
            current_city_adj = input('In one word, describe the city: ').title()
            self.city_dict[current_city_name] = current_city_adj
    # End populate_dict

    def display_dict(self):
        print(self.city_dict)
    # End display_dict

    def reset_dict(self):
        self.city_dict = {}
    # End reset_dict

    def display_as_json(self):
        json_dict = json.dumps(self.city_dict, sort_keys = True)
        print(json_dict)
    # End display_as_json
# End JSONProcessing

pdf_example = PDFProcessing('meetingminutes.pdf')
word_example = WordProcessing('demo.docx')
csv_example = CSVProcessing('example.csv')
json_example = JSONProcessing()

print('All classes have been defined and instantiated.')

#--------------- Requirement #4 ------------------#
'''Determine and display the number of pages in meetingminutes.pdf.'''
print('\nRequirement #4\n')
pdf_example.display_num_pages()

#--------------- Requirement #5 ------------------#
'''Ask the user to enter a page number and display the text on that page.'''
print('\nRequirement #5\n')
pdf_example.display_page_data()
pdf_example.close_pdf()

#--------------- Requirement #6 ------------------#
'''Determine and display the number of paragraphs in demo.docx.'''
print('\nRequirement #6\n')
word_example.display_num_paragraphs()

#--------------- Requirement #7 ------------------#
'''Ask the user to enter a paragraph number and display the text of that paragraph.'''
print('\nRequirement #7\n')
word_example.display_paragraph_text()

#--------------- Requirement #8 ------------------#
'''Display the contents of example.csv.'''
print('\nRequirement #8\n')
csv_example.display_contents()

#--------------- Requirement #9 ------------------#
'''Ask the user to enter data and update example.csv with that data.'''
print('\nRequirement #9\n')
csv_example.write_data()
csv_example.close_csv()

#--------------- Requirement #10 ------------------#
'''Ask the user to enter seven cities and an adjective for each city (e.g. Austin, Awesome).
Enter the data into a Python dictionary.'''
print('\nRequirement #10\n')
json_example.update_dict()

#--------------- Requirement #11 ------------------#
'''Convert the Python dictionary to a string of JSON-formatted data. Display the JSON data.'''
print('\nRequirement #11\n')
json_example.display_as_json()

#--------------- Requirement #12 ------------------#
'''Print a statement explaining your experiences with Program13.
Make this authentic (minimum of 2-3 sentences).'''
print('\nRequirement #12\n')
print('''This was a fun little assignment that broadened my perspective on how robust some file handeling modules can be.
I had tried my hand at JSON and CSV data handeling already, but found it fun to play around with PDFs and Docs as well.
After reading the documentation for these modules I can see that there are so many more things that you can do with them.
I can't wait to play around with them some more and see what I can implement in my future projects.''')